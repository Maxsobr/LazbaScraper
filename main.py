# -*- coding: utf-8 -*-
import os

import pandas as pd
import GenerateComments as gn
import ChartGeneration as chartGeneration
import ending as end
import ploting as plott
from docx import Document
from docx.shared import Pt, RGBColor
from docx.shared import Inches


import re



# Read the CSV file into a DataFrame
df = pd.read_csv('Test.csv')

nameColumn = "Укажи своё имя и фамилию"

RowIndex = 0
# Access a cell value by column name and index
# nameSurname = df.loc[RowIndex, nameColumn]

# Print the cell value
# print(nameSurname)

emailColumn = "Укажи свой E-Mail"

# email = df.loc[RowIndex, emailColumn]



# Print the cell value
# print(email)

firstQuestion = "Если кто-то захочет вас прогнать, вы почувствуете необходимость что-то предпринять в связи с этим?"
# lastQuestion = "Вы любили в детстве играть с оружием?"


# firstAnswer = df.loc[RowIndex, firstQuestion]


# Print the cell value
# print(firstAnswer)

startColumnIndex = df.columns.get_loc(firstQuestion)

# endColumnIndex = df.columns.get_loc(lastQuestion)
#
# i = startColumnIndex
# while i<= endColumnIndex :
#     cell_value = df.iloc[RowIndex, i]
#     print(df.columns[i] + cell_value + "\n")
#     i += 1


# Specify the file path
file_path = 'Calculation rule.xlsx'

# # Loop through each sheet and perform operations
# for sheet_name, sheet_data in all_sheets.items():
#     # Process the data from each sheet
#     print(f"Sheet Name: {sheet_name}")
#     print(sheet_data.head())  # Example: printing the first few rows
#     print('-' * 30)




my_dict = {}
tempEvaluation = 0



def evaluateScore(sheetIndex, file_path, resultSheet,startColumnIndex, ChelRowIndex):

    startColumnIndex -=1



    xl = pd.ExcelFile(file_path)

    sheet_names = xl.sheet_names
    sheet_name = sheet_names[sheetIndex]
    df = pd.read_excel(file_path, sheet_name=sheet_name)
    # print(f"Sheet Name: {sheet_name}")
    # print(df.head())

    totalPoints = 0
    for index, row in df.iterrows():

        if resultSheet.iloc[ChelRowIndex, startColumnIndex + int (row['Вопрос'])] == "Да" and row['Правильный ответ'] == "+":
            totalPoints+=1
        elif resultSheet.iloc[ChelRowIndex, startColumnIndex + int (row['Вопрос'])] == "Нет" and row['Правильный ответ'] == "-":
            totalPoints+=1
        elif resultSheet.iloc[ChelRowIndex, startColumnIndex + int(row['Вопрос'])] == "Не уверен" :
            totalPoints += 0

    my_dict = {}

    # Add elements dynamically
    my_dict[sheet_name] = totalPoints
    return my_dict


        # Access row data using column names or indices
        # print(f"Row {index + 1}: {row['Вопрос']}, {row['Правильный ответ']})  # Example

    # if resultSheet.
    #
    #     cell_value = df.iloc[RowIndex, i]
    #     print(df.columns[i] + cell_value + "\n")

class Chel:
    def __init__(self, name, email, dictResults):
        self.name = name
        self.email = email
        self.dictResults = dictResults

    def __str__(self):
        return f"Name: {self.name}, Email: {self.email}, Results: {self.dictResults}"

    def clean(self):
        newDict = {}
        for key, value in self.dictResults.items():

            if key == "Маскулинность - феминность":
                newDict["Маскулинность / феминность"] = value
            elif key != "Догматизм":
                newDict[key] = value


        self.dictResults = newDict
        # newDict = {}
        # for key, value in self.dictResults[0].items():
        #     if "Догматизм" in self.dictResults:
        #         self.dictResults[0].pop("Догматизм")
        #     if "Маскулинность - феминность" in self.dictResults:
        #         self.dictResults[0]["Маскулинность / феминность"] = self.dictResults[0].pop("Маскулинность - феминность")
        # #
        #     if key != "Догматизм":
        #         newDict[key] = value
        #     elif key == "Маскулинность - феминность":
        #         newDict.pop("Маскулинность - феминность")
        #         newDict["Маскулинность / феминность"] = value
        #
        # self.dictResults = newDict
        #
        # print(newDict)

        # if "Догматизм" in self.dictResults:
        #     del self.dictResults.pop()["Догматизм"]
        # if "Маскулинность - феминность" in self.dictResults:
        #     self.dictResults["Маскулинность / феминность"] = self.dictResults.pop("Маскулинность - феминность")

ChelList = []

for index, row in df.iterrows():
    resultDict = {}
    for sheetIndex in range(7):
        resultDict.update(evaluateScore(sheetIndex, file_path, df, startColumnIndex, index))

    nameSurname = df.loc[index, nameColumn]
    # print(nameSurname, "\n")
    # print(resultDict)
    # print("-" * 30)
    email = df.loc[RowIndex, emailColumn]

    ChelList.append(Chel(nameSurname, email, resultDict))

for i in ChelList:
    i.clean()
    print(i)



# chartGeneration.testChart(ChelList[0])
#
# for i in ChelList:
#     resultString = gn.generateComments(i)
#     end.endString(resultString, i)
#

# Adding a paragraph
#
for i in ChelList:
    resultString = gn.generateComments(i)

    resultString += end.endString(i)

    document = Document()

    filename = plott.boxgraph(i)
    document.add_picture(filename)
    # if
    document.add_paragraph(resultString)
    directory = "resultsDOCX"

    # Create directory if it doesn't exist
    os.makedirs(directory, exist_ok=True)

    output_filename = '{}.docx'.format(
        i.name)  # Example with string formatting

    # Full path to the file
    filepath = os.path.join(directory, output_filename)

    document.save(filepath)





# путь к папке с документами
path = 'resultsDOCX'

search_path = "resultsPng"



# обходим все файлы в папке
for filename in os.listdir(path):
    # проверяем, что это Word-файл
    if filename.endswith('.docx'):
        file_path = os.path.join(path, filename)
        doc = Document(file_path)

        # Создаем новый документ, в который будем записывать отформатированный текст
        new_doc = Document()

        style = new_doc.styles['Normal']
        font = style.font
        font.name = 'Evolventa'
        font.size = Pt(13)

        file2 = os.path.splitext(filename)[
                    0] + '.png'
        print(file2)

        search_file_path = os.path.join(search_path, file2)
        if os.path.exists(search_file_path):
            new_doc.add_picture(search_file_path)


        for para in doc.paragraphs:
            # Разделяем текст на абзацы по двум переходам на новую строчку
            paragraphs = re.split('\n\n|\n \n|\n\n |\n \n | \n \n |  ', para.text)

            for p in paragraphs:
                # Если абзац не пустой, добавляем его в новый документ с отступом
                if p.strip():
                    new_para = new_doc.add_paragraph()
                    # new_para.paragraph_format.first_line_indent = Pt(36)  # Примерно соответствует табуляции.
                    run = new_para.add_run(p)
                    new_para.style = style
                    run.font.color.rgb = RGBColor(255, 255, 255)

        savePath = os.path.join(path, filename)

        # сохраняем изменения
        new_doc.save(savePath)



